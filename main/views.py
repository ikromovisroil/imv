from django.shortcuts import render,redirect,HttpResponseRedirect,get_object_or_404
from unicodedata import category
from .models import *
from django.urls import reverse
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.db.models import Q
from django.db.models import Count, Prefetch
from docx import Document
import os
from django.http import HttpResponse
from django.conf import settings
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor,Inches
from datetime import datetime
from docx.oxml import OxmlElement
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from collections import defaultdict
import json
from django.core.serializers.json import DjangoJSONEncoder
from django.contrib.admin.models import LogEntry
from django.contrib.contenttypes.models import ContentType
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.template.loader import render_to_string
from django.http import JsonResponse

def global_data(request):
    return {
        "global_organizations": Organization.objects.all(),
        "global_categorys": Category.objects.all(),
    }


@login_required
def profil(request):
    context = {
        'employee': Employee.objects.all()
    }
    return render(request, 'main/profil.html', context)


@login_required
def get_employee_files(request):
    emp_id = request.GET.get("employee_id")

    deeds = Deed.objects.filter(receiver_id=emp_id).select_related(
        "sender", "receiver"
    ).order_by("-id")

    html = render_to_string("main/employee_files.html", {"deeds": deeds})
    return JsonResponse({"html": html})


@login_required
def deed_post(request):
    file = request.FILES.get('file')
    message = request.POST.get('message')
    receiver_id = request.POST.get("receiver_id")

    receiver = Employee.objects.filter(id=receiver_id).first()
    sender = Employee.objects.filter(user=request.user).first()

    if not sender:
        return HttpResponse("Sender topilmadi", status=400)

    if not receiver:
        return HttpResponse("Receiver topilmadi", status=400)

    Deed.objects.create(
        sender=sender,
        receiver=receiver,
        file=file,
        message=message,
        status="viewed"
    )

    messages.success(request, "Dalolatnoma yuborildi!")
    return redirect("profil")


@login_required
def index(request):
    organizations = Organization.objects.all()
    categorys = Category.objects.all()

    chart_data = []

    for cat in categorys:
        row = {
            "category": cat.name,   # x o‘qi uchun
        }
        for org in organizations:
            count = Technics.objects.filter(
                employee__organization=org,
                category=cat,
            ).count()
            # JS uchun field: org_1, org_2 ...
            row[f"org_{org.id}"] = count
        chart_data.append(row)

    pie_data = []

    for org in organizations:
        total = Technics.objects.filter(
            employee__organization=org
        ).count()

        pie_data.append({
            "name": org.name,
            "count": total
        })
    organizations1 = Organization.objects.all().annotate(
        technics_count=Count('employee__technics', distinct=True)
    )
    logs = LogEntry.objects.select_related('user', 'content_type').order_by('-action_time')[:10]

    context = {
        "logs": logs,
        "organizations1": organizations1,
        "organizations": organizations,
        "categorys": categorys,
        "chart_data": json.dumps(chart_data, cls=DjangoJSONEncoder),
        "pie_data": json.dumps(pie_data, cls=DjangoJSONEncoder),

    }
    return render(request, "main/index.html", context)


@login_required
def technics(request, pk=None):

    # 1️⃣ CATEGORY FILTER
    if pk:
        category = get_object_or_404(Category, pk=pk)
        technics_qs = Technics.objects.filter(category=category)
    else:
        category = None
        technics_qs = Technics.objects.all()

    # 2️⃣ FILTER PARAMETRLAR
    org_id = request.GET.get('organization')
    dep_id = request.GET.get('department')
    dir_id = request.GET.get('directorate')
    div_id = request.GET.get('division')

    # 3️⃣ FILTERLASH
    if org_id:
        technics_qs = technics_qs.filter(employee__organization_id=org_id)
    if dep_id:
        technics_qs = technics_qs.filter(employee__department_id=dep_id)
    if dir_id:
        technics_qs = technics_qs.filter(employee__directorate_id=dir_id)
    if div_id:
        technics_qs = technics_qs.filter(employee__division_id=div_id)

    # 4️⃣ TEXNIKALAR SONI
    total_count = technics_qs.count()

    # 5️⃣ XODIM BO‘YICHA GURUHLASH
    grouped = defaultdict(list)
    for t in technics_qs.select_related('employee', 'employee__user', 'category').order_by(
            'employee__user__last_name',
            'employee__user__first_name',
            'category__name',
            'name',
    ):
        grouped[t.employee].append(t)

    # 6️⃣ CONTEXT TAYYORLASH
    context = {
        'category': category,
        'grouped_technics': grouped.items(),
        'total_count': total_count,

        # Filter selectlar uchun
        'organizations': Organization.objects.all(),
        'departments': Department.objects.all(),
        'directorate': Directorate.objects.all(),
        'division': Division.objects.all(),

        # Selected qiymatlar
        'selected_org': org_id,
        'selected_dep': dep_id,
        'selected_dir': dir_id,
        'selected_div': div_id,  # ✅ TUZATILGAN
    }

    return render(request, 'main/technics.html', context)


@login_required
def ajax_load_departments(request):
    org_id = request.GET.get('organization')

    if not org_id or org_id == "None":
        return JsonResponse([], safe=False)

    departments = Department.objects.filter(
        organization_id=org_id,
        is_active=True
    ).values('id', 'name')

    return JsonResponse(list(departments), safe=False)


@login_required
def ajax_load_directorate(request):
    dep_id = request.GET.get('department')

    if not dep_id or dep_id == "None":
        return JsonResponse([], safe=False)

    directorate = Directorate.objects.filter(
        department_id=dep_id,
        is_active=True
    ).values('id', 'name')

    return JsonResponse(list(directorate), safe=False)


@login_required
def ajax_load_division(request):
    dir_id = request.GET.get('directorate')

    if not dir_id or dir_id == "None":
        return JsonResponse([], safe=False)

    division = Division.objects.filter(
        directorate_id=dir_id,
        is_active=True
    ).values('id', 'name')

    return JsonResponse(list(division), safe=False)


@login_required
def organization(request, pk):
    organizations = (
        Organization.objects
        .annotate(
            technics_count=Count('employee__technics', distinct=True)
        )
        .prefetch_related(
            'employee_set__technics_set'
        )
        .get(pk=pk)
    )

    departments = (
        Department.objects
        .filter(organization_id=pk)
        .annotate(
            technics_count=Count('employee__technics', distinct=True)
        )
        .prefetch_related('employee_set__technics_set')
    )

    directorates = (
        Directorate.objects
        .filter(department__organization_id=pk)
        .annotate(
            technics_count=Count('employee__technics', distinct=True)
        )
        .prefetch_related('employee_set__technics_set')
    )

    divisions = (
        Division.objects
        .filter(directorate__department__organization_id=pk)
        .annotate(
            technics_count=Count('employee__technics', distinct=True)
        )
        .prefetch_related('employee_set__technics_set')
    )

    context = {
        'organizations': organizations,
        'departments': departments,
        'directorates': directorates,
        'divisions': divisions,
    }
    return render(request, 'main/organization.html', context)


@login_required
def replace_text_in_textboxes(element, replacements):
    """
    DOCX fayldagi barcha text box (shape) ichidagi matnlarni almashtiradi.
    """
    for child in element.iter():
        if child.tag == qn('w:t'):  # w:t -> Word text node
            text = child.text
            if text:
                for old, new in replacements.items():
                    if old in text:
                        child.text = text.replace(old, new)


@login_required
def get_technics_count(request):
    org_id = request.GET.get('org_id')
    dep_id = request.GET.get('dep_id')
    dir_id = request.GET.get('dir_id')
    div_id = request.GET.get('div_id')

    # Asosiy querysetlar (QuerySet sifatida qoldiramiz)
    kompyuterlar = Technics.objects.filter(
        category__name__in=['Kompyuter', 'Planshet', 'Noutbook', 'Doska']
    )
    printerlar = Technics.objects.filter(
        category__name__in=['A4 Printer', 'A3 Printer', 'scaner']
    )

    # Filtrlar to‘g‘ri qo‘llash
    if org_id:
        kompyuterlar = kompyuterlar.filter(employee__organization_id=org_id)
        printerlar = printerlar.filter(employee__organization_id=org_id)

    if dep_id:
        kompyuterlar = kompyuterlar.filter(employee__department_id=dep_id)
        printerlar = printerlar.filter(employee__department_id=dep_id)

    if dir_id:
        kompyuterlar = kompyuterlar.filter(employee__directorate_id=dir_id)
        printerlar = printerlar.filter(employee__directorate_id=dir_id)

    if div_id:
        kompyuterlar = kompyuterlar.filter(employee__division_id=div_id)
        printerlar = printerlar.filter(employee__division_id=div_id)

    # JSON uchun faqat sonlar qaytaramiz
    return JsonResponse({
        "kompyuterlar": kompyuterlar.count(),
        "printerlar": printerlar.count(),
    })


@login_required
def document_get(request):
    """GET so‘rovi uchun sahifani ko‘rsatish"""
    context = {
        'divisions': Division.objects.all(),
        'directorates': Directorate.objects.all(),
        'departments': Department.objects.all(),
        'organizations': Organization.objects.all(),
    }
    return render(request, 'main/document.html', context)


@login_required
def document_post(request):
    """POST so‘rovi uchun dalolatnoma yaratish"""
    oylar = [
        "yanvarda", "fevralda", "martda", "aprelda", "mayda", "iyunda",
        "iyulda", "avgustda", "sentabrda", "oktabrda", "noyabrda", "dekabrda"
    ]

    if request.method != 'POST':
        return redirect('document_get')

    org_id = request.POST.get('organization')
    dep_id = request.POST.get('department')
    dir_id = request.POST.get('directorate')
    div_id = request.POST.get('division')
    post_id = request.POST.get('post_id')
    fio_id = request.POST.get('fio_id')
    date_id = request.POST.get('date_id')
    namber_id = request.POST.get('namber_id')
    rim_id = request.POST.get('rim_id')

    # MODEL OBYEKTLARNI OLISH (bo‘sh bo‘lsa None bo‘ladi)
    org = Organization.objects.filter(id=org_id).first() if org_id else None
    dep = Department.objects.filter(id=dep_id).first() if dep_id else None
    dir = Directorate.objects.filter(id=dir_id).first() if dir_id else None
    div = Division.objects.filter(id=div_id).first() if div_id else None

    # 📅 Sanani formatlash
    formatted_date = ''
    if date_id:
        try:
            dt = datetime.strptime(date_id.strip(), "%Y-%m-%d").date()
            oy_nomi = oylar[dt.month - 1]
            formatted_date = f"{dt.year} yil {dt.day}-{oy_nomi}"
        except:
            formatted_date = str(date_id)

    # ---------- TEXNIKA SANASH QISMI (TUZATILGAN) ----------
    komp_qs = Technics.objects.filter(
        category__name__in=['Kompyuter', 'Planshet', 'Noutbook', 'Doska']
    )
    prin_qs = Technics.objects.filter(
        category__name__in=['A4 Printer', 'A3 Printer', 'scaner']
    )

    # Eng past darajada tanlangan obyekt
    full_name = None

    if div:
        full_name = div
        komp_count = komp_qs.filter(employee__division=div).count()
        prin_count = prin_qs.filter(employee__division=div).count()

    elif dir:
        full_name = dir
        komp_count = komp_qs.filter(employee__directorate=dir).count()
        prin_count = prin_qs.filter(employee__directorate=dir).count()

    elif dep:
        full_name = dep
        komp_count = komp_qs.filter(employee__department=dep).count()
        prin_count = prin_qs.filter(employee__department=dep).count()

    elif org:
        full_name = org
        komp_count = komp_qs.filter(employee__organization=org).count()
        prin_count = prin_qs.filter(employee__organization=org).count()

    else:
        return HttpResponse("Tashkilot / bo‘lim tanlanmagan!", status=400)

    # ---------- DOC SHABLONINI O'ZGARTIRISH ----------
    template_path = os.path.join(settings.MEDIA_ROOT, 'document', 'dalolatnoma1.docx')
    if not os.path.exists(template_path):
        return HttpResponse("Shablon fayl topilmadi!", status=404)

    doc = Document(template_path)

    replacements = {
        'DEPARTMENT': full_name.name,
        'POST': post_id or '',
        'FIO': fio_id or '',
        'DATA': formatted_date or '',
        'NAMBER': namber_id or '',
        'RIM': rim_id or '',
        'STYLE': full_name.name,
        'KOMPCOUNT': str(komp_count),
        'PRINCOUNT': str(prin_count),
    }

    # Matn almashtirish
    for p in doc.paragraphs:
        for run in p.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if old in ['STYLE', 'FIO', 'DATA', 'NAMBER', 'KOMPCOUNT', 'PRINCOUNT']:
                        run.font.bold = True

    # Text boxlarni ham yangilash
    replace_text_in_textboxes(doc.element.body, replacements)

    # Faylni qaytarish
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="dalolatnoma_yangi.docx"'
    doc.save(response)
    return response


@login_required
def ajax_load_technics(request):
    org_id = request.GET.get('organization') or None
    dep_id = request.GET.get('department') or None
    dir_id = request.GET.get('directorate') or None
    div_id = request.GET.get('division') or None

    # Asosiy querysetlar
    kompyuterlar = Technics.objects.filter(
        category__name__in=['Kompyuter', 'Planshet', 'Noutbook', 'Doska']
    )
    printerlar = Technics.objects.filter(
        category__name__in=['A4 Printer', 'A3 Printer', 'scaner']
    )

    if org_id:
        kompyuterlar = kompyuterlar.filter(employee__organization_id=org_id)
        printerlar = printerlar.filter(employee__organization_id=org_id)

    if dep_id:
        kompyuterlar = kompyuterlar.filter(employee__department_id=dep_id)
        printerlar = printerlar.filter(employee__department_id=dep_id)

    if dir_id:
        kompyuterlar = kompyuterlar.filter(employee__directorate_id=dir_id)
        printerlar = printerlar.filter(employee__directorate_id=dir_id)

    if div_id:
        kompyuterlar = kompyuterlar.filter(employee__division_id=div_id)
        printerlar = printerlar.filter(employee__division_id=div_id)

    data_komp = []
    for idx, t in enumerate(kompyuterlar, start=1):
        data_komp.append({
            "num": idx,
            "name": t.name or "",
            "serial": t.serial or "",
            "moc": t.moc or "",
        })

    data_printer = []
    for idx, t in enumerate(printerlar, start=1):
        data_printer.append({
            "num": idx,
            "name": t.name or "",
            "serial": t.serial or "",
        })

    return JsonResponse({
        "kompyuterlar": data_komp,
        "printerlar": data_printer,
    })


@login_required
def hisobot_get(request):
    """GET so‘rovi uchun sahifani ko‘rsatish"""
    context = {
        'divisions': Division.objects.all(),
        'directorates': Directorate.objects.all(),
        'departments': Department.objects.all(),
        'organizations': Organization.objects.all(),
    }
    return render(request, 'main/hisobot.html', context)


@login_required
def set_table_borders(table):
    tbl = table._tbl

    # tblPr mavjud bo'lmasa — yaratamiz
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tblPr)

    # tblBorders yaratish yoki topish
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)

    # 6 ta chiziq
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = borders.find(qn(f'w:{side}'))
        if element is None:
            element = OxmlElement(f'w:{side}')
            borders.append(element)

        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:color'), '000000')
        element.set(qn('w:space'), '0')


@login_required
def set_cell_border(cell, **kwargs):
    """Word jadval hujayrasi atrofida chiziq (border) chizish uchun."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f"w:{edge}")
            for key in ['val', 'sz', 'space', 'color']:
                val = edge_data.get(key)
                if val:
                    element.set(qn(f"w:{key}"), str(val))
            tcPr.append(element)


@login_required
def set_cell_text(cell, text, bold=False, center=False, size=11):
    cell.text = text or ''
    for p in cell.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if center else WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.bold = bold


@login_required
def create_table(doc, title, data, headers):
    """
    Jadvalni widths bo‘yicha yaratadi.
    Agar data bo‘sh bo‘lsa — heading ham, jadval ham chiqmaydi.
    """

    if not data:
        return None, None

    # --- Sarlavha (Heading) ---
    heading = doc.add_paragraph()
    r = heading.add_run(title)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

    # Ustunlar soni
    col_count = len(headers)

    # --- Widths tanlash ---
    if col_count == 4:  # KOMPYUTER JADVALLARI
        widths = [
            Inches(0.5),
            Inches(2),
            Inches(2),
            Inches(2),
        ]

    elif col_count == 3:  # PRINTER JADVALLARI
        widths = [
            Inches(0.5),
            Inches(2.2),
            Inches(2.2),
        ]

    else:
        # default teng taqsimlash (foydalanilmaydi)
        widths = [Inches(6 / col_count)] * col_count

    # --- Jadval yaratish ---
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    # --- Headerlar ---
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].width = widths[i]
        set_cell_text(hdr_cells[i], header, bold=True, center=True)

    # --- Data qatorlari ---
    for idx, item in enumerate(data, start=1):
        row = table.add_row()
        row.height = Inches(0.25)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # Kompyuter jadvali
        if col_count == 4:
            values = [
                f"{idx}.",
                item.name or "",
                item.serial or "",
                getattr(item, "moc", "") or "",
            ]

        # Printer jadvali
        elif col_count == 3:
            values = [
                f"{idx}.",
                item.name or "",
                item.serial or "",
            ]

        # Hujayralarga yozish
        for col_index, val in enumerate(values):
            set_cell_text(
                row.cells[col_index],
                val,
                center=(col_index == 0)  # 1-ustun center
            )
            row.cells[col_index].width = widths[col_index]

    doc.add_paragraph()  # minimal bo‘sh joy
    return heading, table


@login_required
def hisobot_post(request):
    if request.method != 'POST':
        return redirect('hisobot_get')

    # tanlanganlar
    org_id = request.POST.get('organization')
    dep_id = request.POST.get('department')
    dir_id = request.POST.get('directorate')
    div_id = request.POST.get('division')

    org = Organization.objects.filter(id=org_id).first() if org_id else None
    dep = Department.objects.filter(id=dep_id).first() if dep_id else None
    dir = Directorate.objects.filter(id=dir_id).first() if dir_id else None
    div = Division.objects.filter(id=div_id).first() if div_id else None

    # --- Texnikalar ---
    kompyuterlar = Technics.objects.filter(
        category__name__in=['Kompyuter', 'Planshet', 'Noutbook', 'Doska']
    )
    printerlar = Technics.objects.filter(
        category__name__in=['A4 Printer', 'A3 Printer', 'scaner']
    )

    full_name = None
    # *** FILTRLARNI TO‘G‘RI QO‘YISH ***
    if org:
        kompyuterlar = kompyuterlar.filter(employee__organization_id=org_id)
        printerlar = printerlar.filter(employee__organization_id=org_id)
        full_name = org

    if dep:
        kompyuterlar = kompyuterlar.filter(employee__department_id=dep_id)
        printerlar = printerlar.filter(employee__department_id=dep_id)
        full_name = dep

    if dir:
        kompyuterlar = kompyuterlar.filter(employee__directorate_id=dir_id)
        printerlar = printerlar.filter(employee__directorate_id=dir_id)
        full_name = dir

    if div:
        kompyuterlar = kompyuterlar.filter(employee__division_id=div_id)
        printerlar = printerlar.filter(employee__division_id=div_id)
        full_name = div

    # --- DOCX TEMPLATE ---
    template_path = os.path.join(settings.MEDIA_ROOT, 'document', 'dalolatnoma2.docx')
    if not os.path.exists(template_path):
        return HttpResponse("Shablon fayl topilmadi!", status=404)

    doc = Document(template_path)

    # NAME almashtirish
    replacements = {
        'NAME': full_name.name if full_name else '',
    }

    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
                    r.font.bold = True
                    r.font.underline = True

    # TABLE joyini topish
    target_paragraph = None
    for p in doc.paragraphs:
        if 'TABLE' in p.text:
            target_paragraph = p
            p.text = ''
            break

    headers_pc = ['№', 'Rusumi', 'Kompyuter SR:', 'Monitor SR:']
    headers_printer = ['№', 'Rusumi', 'Printer SR:']

    # Kompyuter jadvali
    heading1, table1 = create_table(
        doc,
        "Kompyuterlar (shaxsiy kompyuter, monoblok, noutbuk, planshet va infokioskalar)",
        kompyuterlar,
        headers_pc
    )

    heading2, table2 = create_table(
        doc,
        "Printerlar (lazer, MFU, siyohli va boshqa printerlar)",
        printerlar,
        headers_printer
    )

    # === JOYLASHTIRISH ===
    if target_paragraph:

        # Kompyuterlar bo'lsa
        if table1:
            target_paragraph._p.addnext(heading1._p)
            heading1._p.addnext(table1._tbl)

            # Printerlar ham bo'lsa → kompyuterdan keyin
            if table2:
                table1._tbl.addnext(heading2._p)
                heading2._p.addnext(table2._tbl)

        # Kompyuter bo‘lsa BO‘LMASA, lekin printer BO‘LSA
        elif table2:
            target_paragraph._p.addnext(heading2._p)
            heading2._p.addnext(table2._tbl)

    else:
        # Shablonda TABLE topilmasa
        if table1:
            doc._body._body.append(heading1._p)
            doc._body._body.append(table1._tbl)

        if table2:
            doc._body._body.append(heading2._p)
            doc._body._body.append(table2._tbl)

    # download
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="hisobot_yangi.docx"'
    doc.save(response)
    return response





